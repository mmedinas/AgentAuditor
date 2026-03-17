# -*- coding: utf-8 -*-
import streamlit as st
import os
import pandas as pd
import docx
import json
from io import BytesIO
import altair as alt

# Importando ferramentas da LangChain
from langchain_google_genai import ChatGoogleGenerativeAI
from langchain_core.prompts import ChatPromptTemplate
from langchain_core.output_parsers import StrOutputParser
from langchain_core.messages import HumanMessage, AIMessage

# --- FUNÇÕES DE UTILIDADE (Leitura e Limpeza) ---

def clean_dataframe(df):
    """Mantém apenas colunas úteis para economizar tokens e remove linhas vazias."""
    # Colunas que geralmente importam em engenharia
    palavras_chave = ['item', 'descri', 'especifica', 'qtd', 'quant', 'unid', 'cod', 'part number']
    
    # Converte nomes das colunas para minúsculo para comparar
    df.columns = [str(c).lower() for c in df.columns]
    
    cols_para_manter = [c for c in df.columns if any(p in c for p in palavras_chave)]
    
    if cols_para_manter:
        df = df[cols_para_manter]
    
    return df.dropna(how='all').head(500) # Limite de 500 linhas por aba para segurança

def read_sp_file(file):
    try:
        document = docx.Document(file)
        full_text = [para.text for para in document.paragraphs]
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    full_text.append(cell.text)
        return '\n'.join(full_text)
    except Exception as e:
        st.error(f"Erro ao ler SP: {e}")
        return ""

def read_analysis_files(files):
    all_content = []
    for file in files:
        try:
            file_base_name = os.path.splitext(file.name)[0]
            if file.name.endswith('.csv'):
                df = pd.read_csv(BytesIO(file.getvalue()), sep=None, engine='python')
                df = clean_dataframe(df)
                all_content.append(f"--- LISTA: {file_base_name} ---\n{df.to_string(index=False)}\n")
            elif file.name.endswith('.xlsx'):
                # Lê todas as abas do Excel
                excel_file = pd.ExcelFile(BytesIO(file.getvalue()))
                for sheet_name in excel_file.sheet_names:
                    df = pd.read_excel(excel_file, sheet_name=sheet_name)
                    df = clean_dataframe(df)
                    if not df.empty:
                        all_content.append(f"--- LISTA: {file_base_name} (Aba: {sheet_name}) ---\n{df.to_string(index=False)}\n")
        except Exception as e:
            st.error(f"Erro ao ler arquivo {file.name}: {e}")
    return '\n'.join(all_content)

# --- CONFIGURAÇÃO DOS MODELOS (IA) ---

# Modelo principal para Auditoria (Flash 1.5 - Rápido e Inteligente)
def get_audit_model(api_key):
    return ChatGoogleGenerativeAI(
        model="gemini-1.5-flash",
        google_api_key=api_key,
        temperature=0,
        model_kwargs={"response_mime_type": "application/json"} # Força resposta em JSON
    )

# Modelo para o Chat (Flash 8B - O mais barato/econômico de todos)
def get_chat_model(api_key):
    return ChatGoogleGenerativeAI(
        model="gemini-1.5-flash-8b", 
        google_api_key=api_key,
        temperature=0.1
    )

# --- PROMPTS (Instruções do Sistema) ---

SYSTEM_PROMPT_AUDIT = """Você é um auditor de engenharia. Sua tarefa é comparar a SP (Fonte da Verdade) com as Listas de Materiais.
Responda EXCLUSIVAMENTE no formato JSON abaixo:
{
  "relatorio_markdown": "Seu texto detalhado aqui explicando as divergências...",
  "pendencias": [
    {"Tipo": "FALTANTE", "Lista": "Nome da Lista", "Item": "Nome do Item"},
    {"Tipo": "DISCREPANCIA_TECNICA", "Lista": "Nome da Lista", "Item": "Nome do Item"},
    {"Tipo": "DISCREPANCIA_QUANTIDADE", "Lista": "Nome da Lista", "Item": "Nome do Item"}
  ]
}
"""

SYSTEM_PROMPT_EXTRACT = """Você é um especialista em BOM (Bill of Materials). Extraia os itens da SP.
Responda EXCLUSIVAMENTE no formato JSON:
{
  "relatorio_markdown": "Lista formatada em markdown...",
  "itens": [
    {"Categoria": "Eletrica", "Item": "Cabo X", "Quantidade": "10", "Especificacao": "2.5mm"}
  ]
}
"""

# --- INTERFACE STREAMLIT ---

st.set_page_config(page_title="Agente Auditor v7.0", layout="wide")

# Estilo visual
st.markdown("""<style>
    .stButton>button { width: 100%; border-radius: 5px; }
    .report-box { padding: 20px; border: 1px solid #ddd; border-radius: 10px; background-color: #f9f9f9; }
</style>""", unsafe_allow_html=True)

# Inicialização de estados
if 'chat_history' not in st.session_state: st.session_state.chat_history = []
if 'audit_data' not in st.session_state: st.session_state.audit_data = None
if 'sp_text' not in st.session_state: st.session_state.sp_text = ""
if 'list_text' not in st.session_state: st.session_state.list_text = ""

st.title("🤖 Agente Auditor v7.0")

with st.sidebar:
    st.header("Configurações")
    api_key = os.getenv("GOOGLE_API_KEY")
    if not api_key:
        api_key = st.text_input("Insira sua Google API Key:", type="password")
    
    st.divider()
    sp_file = st.file_uploader("1. Documento Base (SP)", type=["docx"])
    list_files = st.file_uploader("2. Listas de Materiais", type=["xlsx", "csv"], accept_multiple_files=True)
    
    st.divider()
    btn_audit = st.button("🔍 Iniciar Auditoria", type="primary")
    btn_extract = st.button("📋 Extrair Lista Mestra")
    if st.button("🗑️ Limpar Chat"):
        st.session_state.chat_history = []
        st.rerun()

# --- LÓGICA PRINCIPAL ---

if btn_audit and sp_file and list_files and api_key:
    with st.spinner("Auditando documentos..."):
        sp_content = read_sp_file(sp_file)
        list_content = read_analysis_files(list_files)
        st.session_state.sp_text = sp_content
        st.session_state.list_text = list_content
        
        model = get_audit_model(api_key)
        prompt = ChatPromptTemplate.from_messages([
            ("system", SYSTEM_PROMPT_AUDIT),
            ("human", f"SP: {sp_content}\n\nLISTAS: {list_content}")
        ])
        
        try:
            response = (prompt | model | StrOutputParser()).invoke({})
            st.session_state.audit_data = json.loads(response)
        except Exception as e:
            st.error(f"Erro no processamento da IA: {e}")

if btn_extract and sp_file and api_key:
    with st.spinner("Extraindo lista..."):
        sp_content = read_sp_file(sp_file)
        st.session_state.sp_text = sp_content
        
        model = get_audit_model(api_key)
        prompt = ChatPromptTemplate.from_messages([
            ("system", SYSTEM_PROMPT_EXTRACT),
            ("human", f"Documento SP: {sp_content}")
        ])
        
        try:
            response = (prompt | model | StrOutputParser()).invoke({})
            st.session_state.audit_data = json.loads(response)
        except Exception as e:
            st.error(f"Erro na extração: {e}")

# --- EXIBIÇÃO DE RESULTADOS ---

if st.session_state.audit_data:
    data = st.session_state.audit_data
    
    col1, col2 = st.columns([1, 1])
    
    with col1:
        st.subheader("Relatório Detalhado")
        st.markdown(data.get("relatorio_markdown", "Sem relatório disponível."))
        
        # Botão de download do relatório
        st.download_button("Baixar Relatório (.md)", data.get("relatorio_markdown", ""), file_name="relatorio_auditoria.md")

    with col2:
        st.subheader("Visualização de Dados")
        # Se for auditoria, mostra pendências
        if "pendencias" in data:
            df_pendencias = pd.DataFrame(data["pendencias"])
            if not df_pendencias.empty:
                st.dataframe(df_pendencias, use_container_width=True)
                
                # Gráfico
                chart = alt.Chart(df_pendencias).mark_bar().encode(
                    x='count()',
                    y='Tipo',
                    color='Tipo'
                ).properties(height=200)
                st.altair_chart(chart, use_container_width=True)
        
        # Se for extração, mostra itens
        if "itens" in data:
            df_itens = pd.DataFrame(data["itens"])
            st.dataframe(df_itens, use_container_width=True)
            csv = df_itens.to_csv(index=False).encode('utf-8-sig')
            st.download_button("Baixar Tabela CSV", csv, "lista_extraida.csv", "text/csv")

    st.divider()

    # --- CHAT TIRA-DÚVIDAS (MODELO 8B MAIS BARATO) ---
    st.subheader("💬 Chat sobre o Projeto")
    
    for msg in st.session_state.chat_history:
        with st.chat_message(msg["role"]):
            st.markdown(msg["content"])

    if user_input := st.chat_input("Pergunte algo sobre a SP ou as Listas..."):
        st.session_state.chat_history.append({"role": "user", "content": user_input})
        with st.chat_message("user"): st.markdown(user_input)
        
        with st.chat_message("assistant"):
            chat_model = get_chat_model(api_key)
            # Enviamos apenas um resumo/contexto para o chat não ficar caro
            contexto = f"CONTEXTO SP: {st.session_state.sp_text[:5000]}\n\nCONTEXTO LISTAS: {st.session_state.list_text[:5000]}"
            
            prompt_chat = ChatPromptTemplate.from_messages([
                ("system", "Você é um assistente técnico. Responda com base no contexto fornecido."),
                ("human", f"{contexto}\n\nPergunta: {user_input}")
            ])
            
            response = (prompt_chat | chat_model | StrOutputParser()).invoke({})
            st.markdown(response)
            st.session_state.chat_history.append({"role": "assistant", "content": response})
